"""Append a CP-2 / CP-3 addendum to the existing .docx plan.

WHY:  — the original `План_курсового_проекта_WAF.docx` was
generated at  and lists what the project *intended* to deliver.
After CP-3 closed  we have actual numbers: tests passing,
ADRs written, pipeline metrics, threshold calibration, attack-bench
results. This script appends a faithful addendum at the document's end
without touching the original sections — the supervisor can compare
plan vs. fact directly.

Run:
    python scripts/build_plan_addendum.py

Idempotent: looks for a marker heading "Приложение Б" and overwrites
that section if present, otherwise appends a fresh one.

Style: follows the IEML methodology already used in the document —
Times New Roman 14pt, line spacing 1.5, justified text. Tables get
a thin grid border.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

DOCX_PATH = Path(__file__).resolve().parents[1] / "План_курсового_проекта_WAF.docx"

ADDENDUM_HEADING = "Приложение Б. Фактические результаты CP-2 и CP-3"

FONT_NAME = "Times New Roman"
FONT_PT = Pt(14)


def _ieml_paragraph(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm: float = 1.25):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(first_line_cm)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = align
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.size = FONT_PT
        # WHY: Cyrillic + Latin face must be set the same way explicitly,
        # otherwise Word falls back to its own default for Cyrillic.
        r.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        r.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        r.element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)


def _apply_table_borders(tbl):
    """Thin black grid on all sides + inside — IEML requires visible cells."""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)


def _heading(doc: Document, text: str, *, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(13)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run.font.name = FONT_NAME
    return p


def _para(doc: Document, text: str, *, justify: bool = True, first_line: float = 1.25):
    p = doc.add_paragraph(text)
    _ieml_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT,
        first_line_cm=first_line,
    )
    return p


def _table(doc: Document, headers: list[str], rows: list[list[str]]):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # WHY: the base doc may not ship "Table Grid" style. We set the
    #      borders explicitly via _apply_table_borders() instead of
    #      relying on a named style.
    head_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        head_cells[i].text = h
        for p in head_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = FONT_NAME
                r.font.size = Pt(13)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.size = Pt(12)
    _apply_table_borders(tbl)
    return tbl


def _drop_existing_addendum(doc: Document) -> None:
    """If the addendum was appended before, remove its body so we don't
    duplicate. We anchor on the heading text and drop everything from
    that paragraph to the end of the document.
    """
    body = doc.element.body
    found = False
    for p in list(doc.paragraphs):
        if p.text.strip().startswith(ADDENDUM_HEADING):
            found = True
        if found:
            body.remove(p._element)
    # Also drop trailing tables that lived inside the addendum.
    if found:
        for t in list(doc.tables):
            if t._element.getparent() is None:
                continue
            # Tables that come *after* the heading: same parent (body),
            # but the heading is gone; we use a heuristic — drop tables
            # whose first cell text is in the known set of addendum
            # tables. Cheap; idempotent.
            first_cell = t.rows[0].cells[0].text.strip().lower() if t.rows else ""
            if first_cell in {"спринт", "файл", "ключ", "контрольная точка"}:
                t._element.getparent().remove(t._element)


def build() -> None:
    if not DOCX_PATH.exists():
        print(f"ERROR: {DOCX_PATH} missing", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(DOCX_PATH))
    _drop_existing_addendum(doc)

    # New page break before the addendum.
    last_para = doc.add_paragraph()
    last_para.add_run().add_break()  # type: ignore[attr-defined]

    _heading(doc, ADDENDUM_HEADING, level=1)
    _para(
        doc,
        "Раздел добавлен по итогам  проекта. В нём зафиксированы "
        "фактические результаты на момент закрытия контрольных точек CP-2 "
        "(оффлайн-метрики ML) и CP-3 (полный гибрид прошёл стенд атак). "
        "Источник чисел — автоматический прогон тестов и attack-bench "
        "(`bench/run.py`) в репозитории проекта на дату завершения ",
    )

    # ── CP-2 — оффлайн-метрики ML ───────────────────────────────────────
    _heading(doc, "Б.1. CP-2 — оффлайн ML-метрики", level=2)
    _para(
        doc,
        "Закрыт в  с тремя моделями (Logistic Regression, XGBoost, "
        "Isolation Forest), обученными на едином стратифицированном split'е "
        "(test_size=0.2, stratify=y). Для каждой supervised-модели "
        "дополнительно посчитаны mean ± std по 5-fold stratified "
        "cross-validation (precision, recall, F1, ROC-AUC). Полный "
        "EvalReport включает FPR-at-recall-0.99, confusion-matrix, "
        "пороги для recall=0.90 и 0.99.",
    )
    _para(
        doc,
        "Запускалось через `make train` и `make train-register`; артефакты "
        "размещены в `ml/models/<version>/`, метаданные — в Postgres "
        "`ml_models`. Калибровка порога — отдельный модуль "
        "`waf_ml.threshold`  с CLI и JSON-отчётом ROC-trace.",
    )

    _table(
        doc,
        headers=["Спринт", "Артефакт", "Объём кода", "Тесты"],
        rows=[
            ["", "ml/src/waf_ml/{features,train,eval,registry}.py", "≈ 700 строк", "30 (features 9, train 4 + 1 CV, eval 6, registry 5, cicids 6)"],
            ["", "ml/datasets/{synthetic,csic,cicids}.py", "≈ 280 строк", "вкл. в 30 выше"],
            ["", "ml/src/waf_ml/drift.py (PSI + KS)", "≈ 200 строк", "12 (PSI invariants, KS, threshold mapping, CLI)"],
            ["", "ml/src/waf_ml/threshold.py (калибровка)", "≈ 170 строк", "12 (sweep, monotonic, FPR-budget, CLI)"],
            ["ИТОГО CP-2", "пакет ml/", "≈ 1 350 строк", "55 тестов, 100 % проходят"],
        ],
    )

    # ── CP-3 — гибрид + attack bench ────────────────────────────────────
    _heading(doc, "Б.2. CP-3 — стенд атак и блокировка", level=2)
    _para(
        doc,
        "Закрыт в  Online-инференс работает в score-augment "
        "режиме (`/api/v1/ml/inspect`, fail-open). Block-mode вшит в "
        "Lua-subrequest `score.lua` за переменной `ml_block_threshold` "
        "(default 1.0 — annotate-only) с тремя независимыми kill-switch'ами "
        "(ADR-0011): UI-слайдер, env-переменная, переключение flavor'а. "
        "AWS WAF-адаптер опционально пушит блоклист в IPSet (ADR-0012).",
    )
    _para(
        doc,
        "Attack-bench harness `bench/run.py` гонит 100 benign + 100 "
        "malicious запросов (SQLi, XSS, path traversal, RCE, SSRF, "
        "Log4Shell, tooling fingerprints) против целевого URL и "
        "вычисляет FPR / FNR / p50 / p95 / p99 latency. Стаб-сервер в "
        "тестах подтверждает, что арифметика TPR + FNR = 1 совпадает "
        "с per-probe результатами с точностью до float epsilon.",
    )

    _table(
        doc,
        headers=["Контрольная точка", "Что должно быть", "Что фактически"],
        rows=[
            ["Калибровка порога", "θ ≤ 1.0 при FPR ≤ 1 %", "θ выбирается как минимальный θ при FPR ≤ target_fpr; ROC-trace в JSON"],
            ["Block-mode", "Включаемый, безопасно отключаемый", "Lua + ml_block_threshold; 3 kill-switch'а (UI, env, flavor)"],
            ["Attack bench", "FPR ≤ 5 %, FNR ≤ 30 %", "CLI-выход 0/2 по этим бюджетам; harness покрыт 5 unit-тестами"],
            ["AWS WAF", "Опциональный outbound", "boto3 IPSet sync, fail-soft, rate-limit 5 мин (10 unit-тестов)"],
            ["Persist θ ", "ml_config с audit-row", "Postgres ml_config + Alembic 0002; 7 тестов API"],
            ["Drift worker ", "Off-band PSI + KS", "backend/workers/drift_worker.py + 5 тестов на стаб-CH"],
        ],
    )

    # ── Сводка артефактов ──────────────────────────────────────────────
    _heading(doc, "Б.3. Сводка тестового покрытия и ADR", level=2)
    _para(
        doc,
        "На дату закрытия  в репозитории присутствуют четыре "
        "независимых пакета: `backend/` (FastAPI-шлюз с RBAC + audit), "
        "`ml/` (оффлайн-конвейер), `ml-service/` (online-инференс), "
        "`bench/` (attack-bench). Каждый пакет имеет свой `pyproject.toml` "
        "и автономный pytest-набор; запуск всех четырёх занимает "
        "несколько секунд и не требует Postgres / ClickHouse / Redis "
        "(используются in-memory подмены и stub-серверы).",
    )

    _table(
        doc,
        headers=["Пакет", "Кол-во тестов", "Что покрывает"],
        rows=[
            ["backend/", "≈ 75", "auth, RBAC, rules, incidents, audit, metrics, ml inspect/explain/threshold, aws_waf, drift_worker"],
            ["ml/", "55", "features (golden), datasets, train, eval, registry, cicids loader, drift, threshold"],
            ["ml-service/", "21", "score, explain, healthz, cache fail-open, model loader fallback"],
            ["bench/", "5", "арифметика FPR/FNR на стаб-сервере, CLI-отчёт"],
            ["ИТОГО", "≈ 156", "ruff clean во всех четырёх пакетах"],
        ],
    )

    _para(
        doc,
        "Архитектурные решения зафиксированы в виде ADR (Architecture "
        "Decision Records) под `docs/adr/`:",
    )

    _table(
        doc,
        headers=["№", "Заголовок", "Спринт"],
        rows=[
            ["0001", "Tech stack: ModSecurity, ClickHouse, FastAPI, React", "1"],
            ["0002", "Repository abstraction + sessions", "3"],
            ["0003", "React + Vite + ручной CSS", "4"],
            ["0004", "Type generation from OpenAPI", "5"],
            ["0006", "ClickHouse materialized views для метрик", "6"],
            ["0007", "ML pipeline: отдельный пакет, 3-модельный split", "7"],
            ["0008", "Online inference: ml-service, fail-open SLO", "8"],
            ["0009", "Drift detection: PSI + KS, frozen baseline", "9"],
            ["0010", "OpenResty + Lua subrequest as opt-in flavour", "9"],
            ["0011", "ML block-mode: threshold, rollback, kill-switches", "10"],
            ["0012", "AWS WAF adapter: optional, IPSet-only, fail-soft", "10"],
        ],
    )

    # ── Заключительный абзац ───────────────────────────────────────────
    _heading(doc, "Б.4. Соответствие методичке", level=2)
    _para(
        doc,
        "Все пункты ТЗ варианта № 14 (расширенного) реализованы и "
        "подтверждены тестами либо демонстрируемым artefact'ом стенда. "
        "Block-mode остаётся выключенным по умолчанию (annotate-only) — "
        "для production-включения требуется явная калибровка через "
        "`make calibrate` и понижение `ml_block_threshold` через UI или "
        "переменную окружения. Это сознательное решение по ADR-0011: "
        "fail-safe-по-умолчанию.",
    )

    doc.save(str(DOCX_PATH))
    print(f"updated: {DOCX_PATH}")


if __name__ == "__main__":
    build()

    doc.save(str(DOCX_PATH))
    print(f"updated: {DOCX_PATH}")


if __name__ == "__main__":
    build()
